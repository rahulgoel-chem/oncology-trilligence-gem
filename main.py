import requests
import json
import time
import os
import smtplib
from datetime import datetime
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from docx import Document
from fpdf import FPDF

# --- CONFIGURATION ---
SEARCH_QUERY = 'Antibody Drug Conjugate OR ADC OR "CAR-T" OR lymphoma OR leukemia OR "solid tumor"'
EMAIL_SENDER = "theranode@gmail.com"
EMAIL_PASSWORD = "kwutyzqjfddzvtlz" 
EMAIL_RECEIVER = "goel.rahul07@gmail.com"
SNAPSHOT_FILE = "trials_snapshot.json"
CHECK_INTERVAL = 4 * 3600  # 4 Hours

# --- MODULE 1: EXTRACTION & CLEANING ---
def fetch_and_clean_data(query):
    base_url = "https://clinicaltrials.gov/api/v2/studies"
    fields = [
        "protocolSection.identificationModule", "protocolSection.statusModule",
        "protocolSection.sponsorCollaboratorsModule", "protocolSection.descriptionModule",
        "protocolSection.designModule", "protocolSection.armsInterventionsModule",
        "protocolSection.outcomesModule", "protocolSection.conditionsModule",
        "protocolSection.contactsLocationsModule"
    ]
    
    params = {"query.cond": query, "fields": ",".join(fields), "pageSize": 100}
    all_trials = {}
    next_token = None

    while True:
        if next_token: params["pageToken"] = next_token
        try:
            response = requests.get(base_url, params=params)
            response.raise_for_status()
            data = response.json()
            
            for study in data.get("studies", []):
                protocol = study.get("protocolSection", {})
                id_mod = protocol.get("identificationModule", {})
                status_mod = protocol.get("statusModule", {})
                
                nct_id = id_mod.get("nctId")
                if not nct_id: continue

                # Extraction with "NA" for blanks
                all_trials[nct_id] = {
                    "Study Title": id_mod.get("briefTitle", "NA"),
                    "Sponsor": protocol.get("sponsorCollaboratorsModule", {}).get("leadSponsor", {}).get("name", "NA"),
                    "Phases": ", ".join(protocol.get("designModule", {}).get("phases", ["NA"])),
                    "Study Status": status_mod.get("overallStatus", "NA"),
                    "Conditions": ", ".join(protocol.get("conditionsModule", {}).get("conditions", ["NA"])),
                    "Interventions": ", ".join([i.get("name", "NA") for i in protocol.get("armsInterventionsModule", {}).get("interventions", [])]),
                    "Enrollment": status_mod.get("enrollmentStruct", {}).get("count", "NA"),
                    "Primary Completion Date": status_mod.get("primaryCompletionDateStruct", {}).get("date", "NA"),
                    "Completion Date": status_mod.get("completionDateStruct", {}).get("date", "NA"),
                    "Start Date": status_mod.get("startDateStruct", {}).get("date", "NA"),
                    "Collaborators": ", ".join([c.get("name") for c in protocol.get("sponsorCollaboratorsModule", {}).get("collaborators", [])]) or "NA",
                    "Locations": ", ".join(list(set([l.get("facility", "") for l in protocol.get("contactsLocationsModule", {}).get("locations", []) if l.get("facility")]))) or "NA"
                }
            
            next_token = data.get("nextPageToken")
            if not next_token: break
        except Exception as e:
            print(f"Extraction Error: {e}")
            break
    return all_trials

# --- MODULE 2: COMPARISON & REPORT GENERATION ---
def compare_and_report(old_data, new_data):
    fields_to_check = ["Study Status", "Sponsor", "Collaborators", "Phases", "Enrollment", "Start Date", "Primary Completion Date", "Completion Date", "Locations"]
    reports = []
    
    for nct_id, current in new_data.items():
        if nct_id in old_data:
            prev = old_data[nct_id]
            changes = [f"updated '{f}' from '{prev.get(f)}' to '{current.get(f)}'" for f in fields_to_check if str(prev.get(f)) != str(current.get(f))]
            
            if changes:
                reports.append(f"\"{current['Sponsor']}\"'s \"{current['Phases']}\" trial evaluating \"{current['Interventions']}\" in patients with \"{current['Conditions']}\" has {'; '.join(changes)}.")
    return reports

# --- MODULE 3: EXPORT TO DOC & PDF ---
def export_files(report_list):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    files = []
    
    # Word Export
    doc = Document(); doc.add_heading('Clinical Trial Updates', 0)
    for r in report_list: doc.add_paragraph(r, style='List Bullet')
    doc_path = f"Report_{timestamp}.docx"; doc.save(doc_path); files.append(doc_path)
    
    # PDF Export
    pdf = FPDF(); pdf.add_page(); pdf.set_font("Arial", size=11)
    for r in report_list: pdf.multi_cell(0, 10, f"- {r}"); pdf.ln(2)
    pdf_path = f"Report_{timestamp}.pdf"; pdf.output(pdf_path); files.append(pdf_path)
    
    return files

# --- MODULE 4: EMAIL ALERTS ---
def send_email_with_attachments(report_list, file_paths):
    msg = MIMEMultipart()
    msg['Subject'] = "URGENT: Clinical Trial Changes Detected"
    msg['From'] = EMAIL_SENDER
    msg['To'] = EMAIL_RECEIVER
    msg.attach(MIMEText("\n".join(report_list)))

    for path in file_paths:
        with open(path, "rb") as f:
            part = MIMEBase("application", "octet-stream")
            part.set_payload(f.read()); encoders.encode_base64(part)
            part.add_header("Content-Disposition", f"attachment; filename={path}")
            msg.attach(part)

    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
        server.login(EMAIL_SENDER, EMAIL_PASSWORD)
        server.send_message(msg)

# --- MAIN SCHEDULING LOOP ---
def main():
    print("Clinical Trials Monitor Started...")
    while True:
        print(f"Checking for updates: {datetime.now()}")
        
        # 1. Load Old Data
        old_data = {}
        if os.path.exists(SNAPSHOT_FILE):
            with open(SNAPSHOT_FILE, 'r') as f: old_data = json.load(f)
        
        # 2. Fetch New Data
        new_data = fetch_and_clean_data(SEARCH_QUERY)
        
        # 3. Compare
        reports = compare_and_report(old_data, new_data)
        
        if reports:
            print(f"Changes found! Generating reports...")
            paths = export_files(reports)
            send_email_with_attachments(reports, paths)
            print("Alerts sent.")
        else:
            print("No changes found.")

        # 4. Save New Data as Old Data for next cycle
        with open(SNAPSHOT_FILE, 'w') as f: json.dump(new_data, f, indent=4)
        
        print(f"Sleeping for 4 hours...")
        time.sleep(CHECK_INTERVAL)

if __name__ == "__main__":
    main()
